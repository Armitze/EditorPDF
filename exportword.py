"""Exportación PDF -> Word (.docx) con fidelidad visual (estilo LovePDF).

Técnica (la misma que usan los conversores en línea buenos, comprobada abriendo
un .docx de LovePDF: imágenes ancladas detrás del texto + cuadros de texto
flotantes):

1. Por cada página se rasteriza SOLO la capa gráfica (colores de fondo,
   degradados, líneas, fotos…) quitando el texto con redacciones de PyMuPDF, y
   se ancla como imagen a página completa DETRÁS del texto (behindDoc).
2. El texto se reconstruye encima con cuadros de texto flotantes (uno por
   línea, en su posición exacta), conservando fuente, tamaño, color, negrita y
   cursiva. El texto queda seleccionable y editable en Word.

Así el documento se ve prácticamente idéntico al PDF (fondos, tablas, fotos)
sin depender de la detección de tablas, que es lo que suele romper el diseño.
"""
import io
import re
from xml.sax.saxutils import escape

import fitz  # PyMuPDF

# Resolución del fondo rasterizado. 150 dpi equilibra nitidez y peso; el texto
# no pierde calidad porque va aparte, como texto real.
BG_DPI = 150

EMU_PER_PT = 12700

_SUBSET_RE = re.compile(r'^[A-Z]{6}\+')


def _emu(pt):
    return int(round(pt * EMU_PER_PT))


def _font_info(name):
    """(familia, negrita, cursiva) a partir del nombre de fuente del PDF.

    'ABCDEF+Arial-BoldMT' -> ('Arial', True, False). Los pesos que en Windows
    son familias propias se conservan con su nombre ('Arial-Black' -> 'Arial
    Black'); si Word no tiene la fuente sustituye por una parecida y, como
    cada línea va posicionada en su sitio, el diseño general se mantiene.
    """
    raw = _SUBSET_RE.sub('', name or '')
    low = raw.lower()
    if 'glyphless' in low:
        # Fuente sintética de la capa OCR de Tesseract: no existe en Word.
        return ('Arial', False, False)
    bold = 'bold' in low
    italic = 'italic' in low or 'oblique' in low
    base = re.split(r'[-,]', raw)[0].strip()
    for suffix in ('PSMT', 'MT', 'PS'):
        if base.endswith(suffix):
            base = base[: -len(suffix)]
    family = base
    if 'black' in low or 'heavy' in low:
        family, bold = base + ' Black', True
    elif 'semibold' in low or 'demibold' in low or 'demi' in low:
        family, bold = base + ' Semibold', True
    elif 'light' in low:
        family = base + ' Light'
    elif 'medium' in low:
        family = base + ' Medium'
    elif 'condensed' in low or 'narrow' in low:
        family = base + ' Narrow'
    return (family or 'Calibri', bold, italic)


def _span_xml(span):
    """Run de Word para un span del PDF (texto + fuente/tamaño/color/estilo)."""
    text = span.get('text', '')
    if not text:
        return ''
    flags = span.get('flags', 0)
    family, bold, italic = _font_info(span.get('font', ''))
    bold = bold or bool(flags & 16)
    italic = italic or bool(flags & 2)
    color = '%06X' % (span.get('color', 0) & 0xFFFFFF)
    half_pts = max(2, int(round(span.get('size', 11) * 2)))
    family = escape(family)
    props = (
        f'<w:rFonts w:ascii="{family}" w:hAnsi="{family}" w:cs="{family}"/>'
        + ('<w:b/>' if bold else '')
        + ('<w:i/>' if italic else '')
        + f'<w:color w:val="{color}"/>'
        + f'<w:sz w:val="{half_pts}"/><w:szCs w:val="{half_pts}"/>'
    )
    return (f'<w:r><w:rPr>{props}</w:rPr>'
            f'<w:t xml:space="preserve">{escape(text)}</w:t></w:r>')


def _textbox_xml(doc_pr_id, x_pt, y_pt, w_pt, h_pt, runs_xml, line_h_pt):
    """Cuadro de texto flotante (DrawingML) anclado a la página en (x, y)."""
    cx, cy = _emu(w_pt), _emu(h_pt)
    line = max(20, int(round(line_h_pt * 20)))   # veinteavos de punto
    return (
        '<w:drawing>'
        f'<wp:anchor distT="0" distB="0" distL="0" distR="0" simplePos="0"'
        f' relativeHeight="{200 + doc_pr_id}" behindDoc="0" locked="0"'
        ' layoutInCell="1" allowOverlap="1">'
        '<wp:simplePos x="0" y="0"/>'
        f'<wp:positionH relativeFrom="page"><wp:posOffset>{_emu(x_pt)}</wp:posOffset></wp:positionH>'
        f'<wp:positionV relativeFrom="page"><wp:posOffset>{_emu(y_pt)}</wp:posOffset></wp:positionV>'
        f'<wp:extent cx="{cx}" cy="{cy}"/>'
        '<wp:effectExtent l="0" t="0" r="0" b="0"/><wp:wrapNone/>'
        f'<wp:docPr id="{doc_pr_id}" name="Texto {doc_pr_id}"/>'
        '<wp:cNvGraphicFramePr/>'
        '<a:graphic><a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">'
        '<wps:wsp xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">'
        '<wps:cNvSpPr txBox="1"/>'
        f'<wps:spPr><a:xfrm><a:off x="0" y="0"/><a:ext cx="{cx}" cy="{cy}"/></a:xfrm>'
        '<a:prstGeom prst="rect"><a:avLst/></a:prstGeom>'
        '<a:noFill/><a:ln><a:noFill/></a:ln></wps:spPr>'
        '<wps:txbx><w:txbxContent><w:p><w:pPr>'
        f'<w:spacing w:before="0" w:after="0" w:line="{line}" w:lineRule="exact"/>'
        f'</w:pPr>{runs_xml}</w:p></w:txbxContent></wps:txbx>'
        '<wps:bodyPr rot="0" wrap="none" lIns="0" tIns="0" rIns="0" bIns="0"'
        ' anchor="t"><a:noAutofit/></wps:bodyPr>'
        '</wps:wsp></a:graphicData></a:graphic></wp:anchor></w:drawing>'
    )


def _background_xml(doc_pr_id, rid, w_pt, h_pt):
    """Imagen de fondo a página completa, anclada DETRÁS del texto."""
    cx, cy = _emu(w_pt), _emu(h_pt)
    return (
        '<w:drawing>'
        '<wp:anchor distT="0" distB="0" distL="0" distR="0" simplePos="0"'
        ' relativeHeight="1" behindDoc="1" locked="0" layoutInCell="1" allowOverlap="1">'
        '<wp:simplePos x="0" y="0"/>'
        '<wp:positionH relativeFrom="page"><wp:posOffset>0</wp:posOffset></wp:positionH>'
        '<wp:positionV relativeFrom="page"><wp:posOffset>0</wp:posOffset></wp:positionV>'
        f'<wp:extent cx="{cx}" cy="{cy}"/>'
        '<wp:effectExtent l="0" t="0" r="0" b="0"/><wp:wrapNone/>'
        f'<wp:docPr id="{doc_pr_id}" name="Fondo {doc_pr_id}"/>'
        '<wp:cNvGraphicFramePr><a:graphicFrameLocks noChangeAspect="1"/></wp:cNvGraphicFramePr>'
        '<a:graphic><a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">'
        '<pic:pic><pic:nvPicPr>'
        f'<pic:cNvPr id="{doc_pr_id}" name="fondo"/><pic:cNvPicPr/></pic:nvPicPr>'
        f'<pic:blipFill><a:blip r:embed="{rid}"/><a:stretch><a:fillRect/></a:stretch></pic:blipFill>'
        f'<pic:spPr><a:xfrm><a:off x="0" y="0"/><a:ext cx="{cx}" cy="{cy}"/></a:xfrm>'
        '<a:prstGeom prst="rect"><a:avLst/></a:prstGeom></pic:spPr>'
        '</pic:pic></a:graphicData></a:graphic></wp:anchor></w:drawing>'
    )


def _page_background_png(src_doc, index, hide_rects=None):
    """PNG de la página SIN su texto (solo gráficos, imágenes y colores).

    `hide_rects` (en puntos PDF) se pintan de blanco sobre la imagen: en un
    escaneo el texto son píxeles de la foto y las redacciones no lo quitan,
    así que se blanquean los recuadros de las palabras que el OCR reconoció
    para que el texto editable que va encima no salga duplicado.
    """
    tmp = fitz.open()
    tmp.insert_pdf(src_doc, from_page=index, to_page=index)
    page = tmp[0]
    page.add_redact_annot(page.rect)
    page.apply_redactions(images=fitz.PDF_REDACT_IMAGE_NONE,
                          graphics=fitz.PDF_REDACT_LINE_ART_NONE,
                          text=fitz.PDF_REDACT_TEXT_REMOVE)
    pix = page.get_pixmap(dpi=BG_DPI, annots=True)
    if hide_rects:
        sc = BG_DPI / 72.0
        pad = 1.5 * sc
        white = (255,) * pix.n
        for r in hide_rects:
            ir = fitz.IRect(int(r.x0 * sc - pad), int(r.y0 * sc - pad),
                            int(r.x1 * sc + pad) + 1, int(r.y1 * sc + pad) + 1)
            ir = ir & pix.irect
            if not ir.is_empty:
                pix.set_rect(ir, white)
    data = pix.tobytes('png')
    tmp.close()
    return data


def convert(pdf_path, target, indexes=None):
    """Convierte `pdf_path` en un .docx en `target`. `indexes`: páginas 0-based."""
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    from docx.shared import Emu, Pt

    src = fitz.open(pdf_path)
    pages = list(indexes) if indexes is not None else list(range(src.page_count))
    docx = Document()
    ns = nsdecls('w', 'wp', 'a', 'r', 'pic')
    next_id = 1

    for k, i in enumerate(pages):
        page = src[i]
        rect = page.rect
        section = docx.sections[0] if k == 0 else docx.add_section(WD_SECTION.NEW_PAGE)
        section.page_width = Emu(_emu(rect.width))
        section.page_height = Emu(_emu(rect.height))
        for attr in ('left_margin', 'right_margin', 'top_margin', 'bottom_margin',
                     'header_distance', 'footer_distance'):
            setattr(section, attr, Emu(0))

        # Párrafo anfitrión: aloja todos los anclajes (fondo + cuadros) de la
        # página. Letra mínima para que no ocupe sitio visible.
        host = docx.add_paragraph()
        host.paragraph_format.space_before = Pt(0)
        host.paragraph_format.space_after = Pt(0)
        host_run = host.add_run('')
        host_run.font.size = Pt(1)

        blocks = page.get_text('dict').get('blocks', [])
        # Palabras de la capa OCR (fuente GlyphLess de Tesseract): su tinta
        # está en la imagen escaneada, no en la capa de texto, así que hay que
        # blanquearla del fondo para no verla dos veces.
        ocr_rects = [fitz.Rect(s['bbox'])
                     for b in blocks if b.get('type') == 0
                     for l in b.get('lines', []) for s in l.get('spans', [])
                     if 'glyphless' in (s.get('font') or '').lower()
                     and s.get('text', '').strip()]

        # 1) Fondo: la página entera sin texto, detrás de todo.
        rid, _ = docx.part.get_or_add_image(
            io.BytesIO(_page_background_png(src, i, ocr_rects)))
        host._p.append(parse_xml(
            f'<w:r {ns}>{_background_xml(next_id, rid, rect.width, rect.height)}</w:r>'))
        next_id += 1

        # 2) Texto: un cuadro flotante por línea, en su posición exacta.
        for block in blocks:
            if block.get('type') != 0:
                continue
            for line in block.get('lines', []):
                runs = ''.join(_span_xml(s) for s in line.get('spans', []))
                if not runs:
                    continue
                x0, y0, x1, y1 = line['bbox']
                w, h = (x1 - x0), (y1 - y0)
                if w <= 0 or h <= 0:
                    continue
                xml = _textbox_xml(next_id, x0, y0, w + 4, h + 2, runs, h)
                host._p.append(parse_xml(f'<w:r {ns}>{xml}</w:r>'))
                next_id += 1

    src.close()
    docx.save(target)
    return target
