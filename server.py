"""Servidor FastAPI: sirve la interfaz (ui/) y expone la API del editor."""
import base64
import binascii
import csv
import io
import os
import re
import sys
import threading
import time

from fastapi import Depends, FastAPI, HTTPException
from fastapi.responses import FileResponse, Response
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel

from pdfcore import DocumentManager, PasswordRequired, PdfState
from updater import UpdateManager


def base_dir():
    if getattr(sys, 'frozen', False):
        return sys._MEIPASS
    return os.path.dirname(os.path.abspath(__file__))


class WindowService:
    """Puente hacia la ventana pywebview; en modo --server usa alternativas."""

    def __init__(self):
        self.window = None
        self.allow_close = False  # el frontend lo pone en True tras confirmar
        self._maximized = False
        # PDF protegido recibido por línea de comandos antes de que la interfaz
        # existiera: se entrega en /api/docs para que pida la clave al arrancar.
        self.pending_locked = None
        # Serializa los avisos al frontend: dos archivos abiertos a la vez no
        # deben ejecutar __openExternal de forma solapada.
        self._notify_lock = threading.Lock()
        # Estado del arrastre de la barra de título (Aero Snap manual).
        self._drag = None  # dict con offset del cursor y geometría previa

    def bind_drag(self, window):
        """Expone al bridge JS de pywebview las funciones de arrastre y snap.

        Con la ventana «frameless» movemos la ventana nosotros mismos para poder
        detectar cuándo el cursor toca el borde superior de la pantalla y, en ese
        caso, maximizar (comportamiento tipo Aero Snap de Windows).
        """
        window.expose(self.drag_start, self.drag_move, self.drag_end,
                      self.toggle_maximize, self.maximize_window)

    def drag_start(self, cursor_x, cursor_y):
        """Inicia el arrastre. Guarda el desfase del cursor respecto a la ventana.

        cursor_x/cursor_y son coordenadas absolutas de pantalla (screenX/screenY).
        """
        if not self.window:
            return
        # Si arrancamos el arrastre desde maximizado, primero restauramos para
        # poder mover la ventana (como hace Windows al «despegar» del borde).
        was_max = self._maximized
        if was_max:
            self.window.restore()
            self._maximized = False
        if was_max:
            # Al restaurar, recolocamos la ventana bajo el cursor (centrada en X,
            # con la barra bajo el ratón) para que no «salte» lejos del puntero.
            w = self.window.width
            off_x, off_y = int(w * 0.5), 16
            self.window.move(int(cursor_x - off_x), max(0, int(cursor_y - off_y)))
        else:
            off_x = cursor_x - self.window.x
            off_y = cursor_y - self.window.y
        self._drag = {'off_x': off_x, 'off_y': off_y, 'snapped': False}

    def drag_move(self, cursor_x, cursor_y, snap_top):
        """Mueve la ventana siguiendo el cursor y maximiza al tocar el borde.

        snap_top: True cuando el cursor está dentro de la franja superior de la
        pantalla (el frontend, que conoce la resolución, lo decide).
        """
        d = self._drag
        if not d or not self.window:
            return
        if snap_top:
            if not d['snapped']:
                self.window.maximize()
                self._maximized = True
                d['snapped'] = True
            return
        # Salimos de la franja: si estábamos maximizados, restauramos y seguimos.
        if d['snapped']:
            self.window.restore()
            self._maximized = False
            d['snapped'] = False
            # Recentramos la ventana restaurada bajo el cursor.
            d['off_x'] = int(self.window.width * 0.5)
            d['off_y'] = 16
        self.window.move(int(cursor_x - d['off_x']), int(cursor_y - d['off_y']))

    def drag_end(self):
        self._drag = None

    def toggle_maximize(self):
        if not self.window:
            return
        if self._maximized:
            self.window.restore()
            self._maximized = False
        else:
            self.window.maximize()
            self._maximized = True

    def maximize_window(self):
        """Maximiza la ventana (idempotente). Usado al abrir un PDF nuevo."""
        if not self.window or self._maximized:
            return
        self.window.maximize()
        self._maximized = True

    def notify_open(self, doc_ids):
        """Avisa al frontend de pestañas nuevas sin ocupar al worker de uvicorn.

        `evaluate_js` es síncrono: espera a que el hilo de la UI ejecute el JS, y
        ese JS (`__openExternal`) vuelve a llamar a la API HTTP. Lanzarlo desde el
        worker que atiende /api/open-external lo dejaría esperando a la UI mientras
        atiende la petición, así que lo hacemos en un hilo aparte y respondemos ya.
        El lock evita que dos archivos abiertos a la vez ejecuten `__openExternal`
        de forma solapada.
        """
        if self.window is None or not doc_ids:
            return

        def run():
            with self._notify_lock:
                try:
                    self.window.evaluate_js(
                        f'window.__openExternal && window.__openExternal({list(doc_ids)!r})')
                except Exception:
                    pass
                self.focus()

        threading.Thread(target=run, daemon=True).start()

    def notify_locked(self, paths):
        """Avisa al frontend de PDFs protegidos para que pida la contraseña.

        Mismo patrón que notify_open: en un hilo aparte, porque `evaluate_js`
        espera al hilo de la UI y el JS vuelve a llamar a la API.
        """
        if self.window is None or not paths:
            return

        def run():
            with self._notify_lock:
                try:
                    self.window.evaluate_js(
                        f'window.__openLocked && window.__openLocked({list(paths)!r})')
                except Exception:
                    pass
                self.focus()

        threading.Thread(target=run, daemon=True).start()

    def focus(self):
        """Trae la ventana al frente (al recibir un archivo de otra instancia).

        Todo el trabajo se hace DENTRO del hilo de la UI y de forma asíncrona.
        Windows exige que las propiedades de una ventana se toquen en su hilo
        propietario: hacerlo desde otro hilo (como hacía `self.window.on_top = True`)
        bloquea para siempre esperando la bomba de mensajes y congela la app —
        justo lo que pasaba al abrir un segundo PDF con la ventana ya abierta.
        `BeginInvoke` no espera respuesta del hilo de UI, así que no puede colgarse.
        """
        if not self.window:
            return
        try:
            from System import Action
            from webview.platforms import winforms as P

            inst = P.BrowserView.instances.get(self.window.uid)
            if inst is None:
                return

            def on_ui():
                # Ya estamos en el hilo de UI: aquí sí es legal tocar la ventana.
                inst.TopMost = True
                inst.Activate()
                inst.TopMost = False

            inst.BeginInvoke(Action(on_ui))
        except Exception:
            # Sin ventana WinForms (modo --server, otras plataformas): no hay nada
            # que traer al frente.
            pass

    def _downloads(self, suggested):
        folder = os.path.join(os.path.expanduser('~'), 'Downloads')
        os.makedirs(folder, exist_ok=True)
        return os.path.join(folder, suggested)

    def open_pdf_dialog(self):
        if not self.window:
            return None
        import webview
        result = self.window.create_file_dialog(
            webview.OPEN_DIALOG, file_types=('Documentos PDF (*.pdf)',))
        return result[0] if result else None

    def open_pdf_dialog_multi(self):
        """Diálogo de apertura con selección múltiple. Devuelve lista de rutas."""
        if not self.window:
            return None
        import webview
        result = self.window.create_file_dialog(
            webview.OPEN_DIALOG, allow_multiple=True,
            file_types=('Documentos PDF (*.pdf)',))
        return list(result) if result else None

    def save_dialog(self, suggested, file_type):
        if not self.window:
            return self._downloads(suggested)
        import webview
        result = self.window.create_file_dialog(
            webview.SAVE_DIALOG, save_filename=suggested, file_types=(file_type,))
        if not result:
            return None
        return result if isinstance(result, str) else result[0]

    def folder_dialog(self):
        if not self.window:
            return os.path.join(os.path.expanduser('~'), 'Downloads')
        import webview
        result = self.window.create_file_dialog(webview.FOLDER_DIALOG)
        if not result:
            return None
        return result if isinstance(result, str) else result[0]

    def control(self, action):
        if not self.window:
            return False
        if action == 'minimize':
            self.window.minimize()
        elif action == 'maximize':
            self.toggle_maximize()
        elif action == 'close':
            # El frontend ya confirmó (o no había cambios): permitir el cierre real.
            self.allow_close = True
            self.window.destroy()
        return True


class OpenBody(BaseModel):
    path: str | list[str] | None = None
    password: str | None = None


class AnnotBody(BaseModel):
    page: int
    kind: str
    color: str = '#f2d024'
    width: float = 3
    opacity: float = 0.8
    rect: list[float] | None = None
    p1: list[float] | None = None
    p2: list[float] | None = None
    text: str = ''
    quads: list[list[float]] | None = None


class PagesBody(BaseModel):
    action: str
    page: int
    scope: str = 'pg'   # pg = solo la página indicada | doc = todo el documento


class SaveBody(BaseModel):
    saveAs: bool = False


class SplitGroup(BaseModel):
    name: str = ''
    range: str = ''


class SplitGroupsBody(BaseModel):
    groups: list[SplitGroup]


class ExportBody(BaseModel):
    fmt: str
    range: str = ''


class WindowBody(BaseModel):
    action: str


class MergeBody(BaseModel):
    mode: str = 'current'  # 'current' = sobre el actual | 'new' = documento nuevo


class SignBody(BaseModel):
    page: int
    rect: list[float]          # [x0, y0, x1, y1] en puntos PDF
    image: str                 # PNG en base64 (con o sin prefijo data:)
    keepRatio: bool = True
    opacity: float = 1.0       # 0..1


class TableExportBody(BaseModel):
    page: int
    fmt: str  # csv | excel


class OcrBody(BaseModel):
    scope: str = 'doc'          # doc | pg
    page: int = 0
    lang: str | None = None     # None = mejor idioma disponible (español si está)
    apply: bool = False         # añadir capa de texto buscable al escaneo
    force: bool = False         # OCR incluso si la página ya tiene texto digital


def _doc_html(pdf: PdfState, indexes):
    body = pdf.html(indexes)
    name = (pdf.info().get('name') or 'documento').rsplit('.', 1)[0]
    return name, f'<!DOCTYPE html><html lang="es"><head><meta charset="utf-8"><title>{name}</title></head><body>{body}</body></html>'


def _tables_html(tables):
    parts = []
    for t in tables:
        rows = ''.join(
            '<tr>' + ''.join(f'<td>{cell}</td>' for cell in row) + '</tr>'
            for row in t['rows'])
        parts.append(f'<table border="1">{rows}</table><br>')
    return ''.join(parts)


_NUM_RE = re.compile(r'^-?\$?[\d.,]*\d$')


def _xlsx_value(text):
    """Celda para Excel: número de verdad si el texto lo es, si no el texto.

    Entiende separador de miles y decimal en ambos convenios (1,234.56 y
    1.234,56): el último . o , con uno o dos dígitos detrás es el decimal.
    Los códigos con cero inicial ("017845") se conservan como texto.
    """
    t = text.strip()
    if not _NUM_RE.match(t) or t in ('-', '$'):
        return text
    s = t.lstrip('-$')
    if s.startswith('0') and '.' not in s and ',' not in s and len(s) > 1:
        return text                      # "017845": código, no cantidad
    last = max(s.rfind('.'), s.rfind(','))
    if last != -1 and 1 <= len(s) - last - 1 <= 2:
        s = s[:last].replace('.', '').replace(',', '') + '.' + s[last + 1:]
    else:
        s = s.replace('.', '').replace(',', '')
    try:
        value = float(s)
    except ValueError:
        return text
    if t.startswith('-'):
        value = -value
    return int(value) if value == int(value) and abs(value) < 1e15 else value


def _write_xlsx(target, tables):
    """Escribe las tablas en un .xlsx real (una hoja, tablas separadas)."""
    from openpyxl import Workbook
    wb = Workbook()
    ws = wb.active
    ws.title = 'Tablas'
    first = True
    for t in tables:
        if not first:
            ws.append([])
        first = False
        for row in t['rows']:
            ws.append([_xlsx_value(c) for c in row])
    wb.save(target)


def _have_openpyxl():
    try:
        import openpyxl  # noqa: F401
        return True
    except ImportError:
        return False


def create_app(manager: DocumentManager, windows: WindowService) -> FastAPI:
    app = FastAPI(title='PDF Editor Pro')
    updates = UpdateManager()
    # Comprobar si hay versión nueva nada más arrancar (en segundo plano: no
    # retrasa el inicio ni molesta si no hay conexión).
    updates.check_async()

    def get_doc(docId: str) -> PdfState:
        """Dependencia: resuelve el PdfState de la pestaña indicada por docId."""
        state = manager.get(docId)
        if state is None:
            raise HTTPException(404, f'La pestaña «{docId}» no existe.')
        return state

    @app.get('/api/ping')
    def ping():
        """Sonda de instancia única: confirma que esta es la ventana primaria."""
        return {'app': 'PDFEditorPro', 'ok': True}

    @app.post('/api/open-external')
    def open_external(body: OpenBody):
        """Abre un PDF llegado de OTRA instancia (doble-clic con la app ya abierta).

        Crea la pestaña en el backend y avisa al frontend para que la muestre y
        traiga la ventana al frente. Devuelve el docId abierto.
        """
        paths = body.path
        if not paths:
            return {'cancelled': True}
        if isinstance(paths, str):
            paths = [paths]
        opened = []
        locked = []      # PDFs con contraseña: los pedirá el frontend
        for path in paths:
            if not os.path.isfile(path):
                continue
            try:
                opened.append(manager.open(path))
            except PasswordRequired:
                # No se descarta en silencio (antes se perdía aquí y el archivo
                # simplemente no aparecía): se avisa al frontend para que pida
                # la clave.
                locked.append(path)
            except Exception:
                pass
        if not opened and not locked:
            raise HTTPException(400, 'No se pudo abrir el archivo recibido.')
        if locked:
            windows.notify_locked(locked)
        if not opened:
            # Todo lo recibido está protegido: la ventana ya está pidiendo la clave.
            windows.focus()
            return {'docs': [], 'locked': locked}
        ids = [d['docId'] for d in opened]
        # Avisar al frontend: que cargue las pestañas nuevas y active la primera.
        # En un hilo aparte: evaluate_js bloquea hasta que el JS termina, y ese JS
        # vuelve a llamar a esta API. Hacerlo aquí colgaría el worker de uvicorn y
        # la ventana (deadlock al abrir dos PDFs seguidos).
        windows.notify_open(ids)
        return {'docs': opened, 'locked': locked}

    @app.get('/api/docs')
    def docs_list():
        """Lista todas las pestañas abiertas (para reconstruir la UI al arrancar).

        `locked` trae el PDF protegido que llegó por línea de comandos, si lo
        hubo: la interfaz lo recoge al arrancar y pide la contraseña. Se entrega
        una sola vez.
        """
        locked = windows.pending_locked
        windows.pending_locked = None
        return {'docs': manager.list(), 'locked': locked}

    @app.get('/api/doc')
    def doc_info(pdf: PdfState = Depends(get_doc)):
        return pdf.info()

    @app.post('/api/open')
    def open_doc(body: OpenBody):
        """Abre uno o varios PDFs, cada uno en su propia pestaña nueva."""
        paths = body.path if body.path else windows.open_pdf_dialog_multi()
        if not paths:
            return {'cancelled': True}
        if isinstance(paths, str):
            paths = [paths]
        opened = []
        for path in paths:
            if not os.path.isfile(path):
                raise HTTPException(404, f'No existe el archivo: {path}')
            try:
                opened.append(manager.open(path, body.password))
            except PasswordRequired as e:
                # 401 = «necesito la contraseña». El frontend la pide y reintenta
                # mandando `password`. Se devuelve la ruta pendiente para que sepa
                # con cuál reintentar cuando se abren varios archivos a la vez.
                raise HTTPException(401, {
                    'message': str(e), 'wrong': e.wrong,
                    'path': path, 'name': os.path.basename(path)})
            except Exception as e:
                raise HTTPException(400, str(e))
        # 'doc' = primera pestaña abierta (la que se activará); 'docs' = todas.
        return {'doc': opened[0], 'docs': opened}

    @app.post('/api/close')
    def close_doc(pdf: PdfState = Depends(get_doc), docId: str = ''):
        manager.close(docId)
        return {'closed': docId, 'docs': manager.list()}

    @app.get('/api/page/{index}')
    def page_png(index: int, scale: float = 2.0, pdf: PdfState = Depends(get_doc)):
        # Mínimo 0.2 (no 1.0): con el zoom alejado el frontend pide la escala
        # EXACTA de pantalla; forzarla a 1 obligaba al navegador a reducir la
        # imagen y el texto se veía suavizado.
        try:
            data = pdf.page_png(index, max(0.2, min(6.0, scale)))
        except Exception as e:
            raise HTTPException(400, str(e))
        return Response(data, media_type='image/png')

    @app.get('/api/thumb/{index}')
    def thumb_png(index: int, dpr: float = 1.0, pdf: PdfState = Depends(get_doc)):
        # `dpr`: densidad de la pantalla (escalado de Windows), para que las
        # miniaturas también se rendericen a los píxeles físicos reales.
        try:
            data = pdf.thumb_png(index, dpr)
        except Exception as e:
            raise HTTPException(400, str(e))
        return Response(data, media_type='image/png')

    @app.get('/api/pagesize/{index}')
    def page_size(index: int, pdf: PdfState = Depends(get_doc)):
        try:
            return pdf.page_size(index)
        except Exception as e:
            raise HTTPException(400, str(e))

    @app.get('/api/pagesizes')
    def page_sizes(pdf: PdfState = Depends(get_doc)):
        """Tamaños de todas las páginas (la vista continua los pide de una vez)."""
        try:
            return pdf.page_sizes()
        except Exception as e:
            raise HTTPException(400, str(e))

    @app.get('/api/words/{index}')
    def words(index: int, pdf: PdfState = Depends(get_doc)):
        try:
            return pdf.words(index)
        except Exception as e:
            raise HTTPException(400, str(e))

    @app.post('/api/annot')
    def add_annot(body: AnnotBody, pdf: PdfState = Depends(get_doc)):
        try:
            return pdf.add_annot(body.page, body.kind, body.color, body.width,
                                 body.opacity, body.rect, body.p1, body.p2, body.text,
                                 body.quads)
        except Exception as e:
            raise HTTPException(400, str(e))

    @app.post('/api/sign')
    def sign(body: SignBody, pdf: PdfState = Depends(get_doc)):
        # El PNG llega en base64; admite el prefijo "data:image/png;base64,".
        raw = body.image.split(',', 1)[-1] if ',' in body.image else body.image
        try:
            png = base64.b64decode(raw, validate=True)
        except (binascii.Error, ValueError):
            raise HTTPException(400, 'La imagen de la firma no es válida.')
        try:
            return pdf.add_signature(body.page, body.rect, png, body.keepRatio,
                                     body.opacity)
        except Exception as e:
            raise HTTPException(400, str(e))

    @app.post('/api/pages')
    def pages(body: PagesBody, pdf: PdfState = Depends(get_doc)):
        try:
            if body.action == 'add':
                return pdf.add_page(body.page)
            if body.action == 'duplicate':
                return pdf.duplicate_page(body.page)
            if body.action == 'delete':
                return pdf.delete_page(body.page)
            # Girar: `scope='doc'` gira todo el documento; si no, la página dada.
            if body.action in ('rotate-left', 'rotate-right'):
                delta = -90 if body.action == 'rotate-left' else 90
                index = None if body.scope == 'doc' else body.page
                return pdf.rotate_page(index, delta)
            raise ValueError(f'Acción desconocida: {body.action}')
        except HTTPException:
            raise
        except Exception as e:
            raise HTTPException(400, str(e))

    @app.post('/api/save')
    def save(body: SaveBody, pdf: PdfState = Depends(get_doc)):
        info = pdf.info()
        if not info.get('open'):
            raise HTTPException(400, 'No hay ningún documento abierto.')
        try:
            if body.saveAs or not info.get('path'):
                target = windows.save_dialog(info.get('name') or 'documento.pdf',
                                             'Documento PDF (*.pdf)')
                if not target:
                    return {'cancelled': True}
                result = pdf.save(target)
            else:
                result = pdf.save()
            result['savedTo'] = pdf.path
            return result
        except Exception as e:
            raise HTTPException(400, str(e))

    @app.post('/api/remove-password')
    def remove_password(body: SaveBody | None = None,
                        pdf: PdfState = Depends(get_doc)):
        """Quita la contraseña del documento y lo guarda sin cifrar.

        Por defecto pregunta dónde guardarlo (no se pisa el original sin querer);
        con saveAs=False reescribe el archivo actual.
        """
        info = pdf.info()
        if not info.get('open'):
            raise HTTPException(400, 'No hay ningún documento abierto.')
        if not info.get('encrypted'):
            raise HTTPException(400, 'El documento no tiene contraseña.')
        target = None
        if not (body and body.saveAs is False):
            name = (info.get('name') or 'documento.pdf')
            base = name[:-4] if name.lower().endswith('.pdf') else name
            target = windows.save_dialog(f'{base}_sin_clave.pdf',
                                         'Documento PDF (*.pdf)')
            if not target:
                return {'cancelled': True}
        try:
            result = pdf.remove_password(target)
        except Exception as e:
            raise HTTPException(400, str(e))
        result['savedTo'] = pdf.path
        return result

    @app.post('/api/split-groups')
    def split_groups(body: SplitGroupsBody, pdf: PdfState = Depends(get_doc)):
        info = pdf.info()
        if not info.get('open'):
            raise HTTPException(400, 'No hay ningún documento abierto.')
        if not body.groups:
            raise HTTPException(400, 'Define al menos un grupo.')
        folder = windows.folder_dialog()
        if not folder:
            return {'cancelled': True}
        try:
            saved = pdf.split_groups([g.model_dump() for g in body.groups], folder)
        except Exception as e:
            raise HTTPException(400, str(e))
        return {'paths': saved, 'folder': folder}

    @app.post('/api/undo')
    def undo(pdf: PdfState = Depends(get_doc)):
        try:
            return pdf.undo()
        except Exception as e:
            raise HTTPException(400, str(e))

    @app.post('/api/redo')
    def redo(pdf: PdfState = Depends(get_doc)):
        try:
            return pdf.redo()
        except Exception as e:
            raise HTTPException(400, str(e))

    @app.post('/api/merge')
    def merge(body: MergeBody | None = None, pdf: PdfState = Depends(get_doc)):
        others = windows.open_pdf_dialog_multi()
        if not others:
            return {'cancelled': True}
        mode = (body.mode if body else 'current')
        try:
            if mode == 'new':
                return pdf.merge_new(others)
            return pdf.merge(others)
        except Exception as e:
            raise HTTPException(400, str(e))

    @app.get('/api/text')
    def text(page: int = 0, scope: str = 'pg', pdf: PdfState = Depends(get_doc)):
        try:
            return {'text': pdf.text(None if scope == 'doc' else page)}
        except Exception as e:
            raise HTTPException(400, str(e))

    @app.get('/api/search')
    def search(q: str = '', pdf: PdfState = Depends(get_doc)):
        """Busca texto en el documento. Devuelve las coincidencias con su página."""
        try:
            hits = pdf.search(q)
        except Exception as e:
            raise HTTPException(400, str(e))
        return {'query': q, 'hits': hits, 'count': len(hits)}

    @app.get('/api/ocr-status')
    def ocr_status():
        """Si el OCR está disponible y con qué idiomas (la interfaz lo consulta)."""
        import ocr as ocr_mod
        return ocr_mod.status()

    @app.post('/api/ocr')
    def run_ocr(body: OcrBody, pdf: PdfState = Depends(get_doc)):
        """Extrae texto con OCR y, si se pide, deja el escaneo buscable.

        scope: 'doc' = todo el documento | 'pg' = solo la página indicada.
        apply: añade la capa de texto invisible sobre las páginas escaneadas.
        """
        import ocr as ocr_mod
        if not ocr_mod.available():
            raise HTTPException(503, 'El OCR no está disponible: falta Tesseract.')
        index = None if body.scope == 'doc' else body.page
        try:
            result = pdf.ocr_text(index, body.lang, body.force)
            if body.apply:
                info = pdf.ocr_apply(index, body.lang)
                result['doc'] = info
        except Exception as e:
            raise HTTPException(400, str(e))
        return result

    @app.get('/api/tables')
    def tables(page: int = 0, pdf: PdfState = Depends(get_doc)):
        try:
            return {'tables': pdf.tables(page)}
        except Exception as e:
            raise HTTPException(400, str(e))

    @app.post('/api/export-table')
    def export_table(body: TableExportBody, pdf: PdfState = Depends(get_doc)):
        try:
            found = pdf.tables(body.page)
        except Exception as e:
            raise HTTPException(400, str(e))
        if not found:
            raise HTTPException(400, 'No se detectaron tablas en la página.')
        name = (pdf.info().get('name') or 'documento').rsplit('.', 1)[0]
        if body.fmt == 'csv':
            target = windows.save_dialog(f'{name}_tabla.csv', 'CSV (*.csv)')
            if not target:
                return {'cancelled': True}
            buf = io.StringIO()
            writer = csv.writer(buf)
            for t in found:
                writer.writerows(t['rows'])
            with open(target, 'w', newline='', encoding='utf-8-sig') as fh:
                fh.write(buf.getvalue())
        elif _have_openpyxl():
            target = windows.save_dialog(f'{name}_tabla.xlsx',
                                         'Libro de Excel (*.xlsx)')
            if not target:
                return {'cancelled': True}
            _write_xlsx(target, found)
        else:
            target = windows.save_dialog(f'{name}_tabla.xls', 'Excel (*.xls)')
            if not target:
                return {'cancelled': True}
            with open(target, 'w', encoding='utf-8') as fh:
                fh.write('<html><head><meta charset="utf-8"></head><body>'
                         + _tables_html(found) + '</body></html>')
        return {'path': target}

    @app.post('/api/export')
    def export(body: ExportBody, pdf: PdfState = Depends(get_doc)):
        info = pdf.info()
        if not info.get('open'):
            raise HTTPException(400, 'No hay ningún documento abierto.')
        try:
            indexes = pdf.parse_range(body.range)
        except Exception:
            raise HTTPException(400, 'Rango de páginas no válido.')
        if not indexes:
            raise HTTPException(400, 'El rango no incluye ninguna página.')
        try:
            if body.fmt == 'word':
                name = (pdf.info().get('name') or 'documento').rsplit('.', 1)[0]
                # .docx REAL reconstruyendo el diseño (exportword, requiere
                # python-docx). Solo si la librería no está instalada se cae al
                # HTML .doc de siempre; un error DE conversión sí se reporta al
                # usuario (no se degrada en silencio a un archivo peor).
                try:
                    import docx  # noqa: F401 — solo comprobar que existe
                    have_docx = True
                except ImportError:
                    have_docx = False
                if have_docx:
                    target = windows.save_dialog(f'{name}.docx',
                                                 'Documento de Word (*.docx)')
                    if not target:
                        return {'cancelled': True}
                    pdf.export_docx(target, indexes)
                else:
                    target = windows.save_dialog(f'{name}.doc',
                                                 'Documento de Word (*.doc)')
                    if not target:
                        return {'cancelled': True}
                    _, html = _doc_html(pdf, indexes)
                    with open(target, 'w', encoding='utf-8') as fh:
                        fh.write(html)
                return {'path': target}
            name, html = _doc_html(pdf, indexes)
            if body.fmt == 'excel':
                found = []
                for i in indexes:
                    found.extend(pdf.tables(i))
                if not found:
                    raise HTTPException(400, 'No se detectaron tablas en el rango indicado.')
                if _have_openpyxl():
                    target = windows.save_dialog(f'{name}.xlsx',
                                                 'Libro de Excel (*.xlsx)')
                    if not target:
                        return {'cancelled': True}
                    _write_xlsx(target, found)
                else:
                    target = windows.save_dialog(f'{name}.xls', 'Excel (*.xls)')
                    if not target:
                        return {'cancelled': True}
                    with open(target, 'w', encoding='utf-8') as fh:
                        fh.write('<html><head><meta charset="utf-8"></head><body>'
                                 + _tables_html(found) + '</body></html>')
            else:
                target = windows.save_dialog(f'{name}.html', 'Página web (*.html)')
                if not target:
                    return {'cancelled': True}
                with open(target, 'w', encoding='utf-8') as fh:
                    fh.write(html)
        except HTTPException:
            raise
        except Exception as e:
            raise HTTPException(400, str(e))
        return {'path': target}

    @app.post('/api/window')
    def window_control(body: WindowBody):
        return {'ok': windows.control(body.action)}

    # ---------- actualización automática ----------
    @app.get('/api/update/status')
    def update_status():
        """Estado del updater: versión actual, si hay una nueva, y progreso."""
        return updates.status()

    @app.post('/api/update/check')
    def update_check():
        """Fuerza una comprobación de versión (además de la del arranque)."""
        updates.check_async()
        return {'ok': True}

    @app.post('/api/update/download')
    def update_download():
        """Empieza a descargar la versión disponible."""
        if not updates.start_download():
            raise HTTPException(400, 'No hay ninguna actualización lista para descargar.')
        return {'ok': True}

    @app.post('/api/update/apply')
    def update_apply():
        """Aplica la actualización descargada y cierra la app para reinstalarla.

        El aplicador externo espera a que este proceso muera, reemplaza la
        carpeta de instalación y vuelve a abrir la app ya actualizada.
        """
        def close_app():
            windows.allow_close = True
            if windows.window:
                try:
                    windows.window.destroy()
                except Exception:
                    os._exit(0)
            else:
                os._exit(0)

        if not updates.apply(close_app):
            st = updates.status()
            raise HTTPException(400, st.get('error') or 'No se pudo aplicar la actualización.')
        return {'ok': True}

    @app.post('/api/update/cancel')
    def update_cancel():
        updates.cancel()
        return {'ok': True}

    app.mount('/', StaticFiles(directory=os.path.join(base_dir(), 'ui'), html=True), name='ui')
    return app
